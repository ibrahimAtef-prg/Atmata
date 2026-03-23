"""
md_to_docx.py — Markdown → Word Document Converter
====================================================
Converts an Aurora-generated Markdown report into a styled .docx file.

Usage:
    python md_to_docx.py --input report.md --output report.docx
"""

from __future__ import annotations
import argparse, os, re, sys
from typing import List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Colour palette (matches Aurora dark theme as close as docx allows) ──────
PURPLE      = RGBColor(0x8B, 0x5C, 0xF6)   # headings
DARK_BG     = RGBColor(0x1E, 0x1B, 0x2E)   # not used on text but kept for ref
BODY_COLOUR = RGBColor(0x1F, 0x29, 0x37)   # near-black body text
GREY        = RGBColor(0x6B, 0x72, 0x80)   # meta / caption
GREEN       = RGBColor(0x10, 0xB9, 0x81)
RED         = RGBColor(0xEF, 0x44, 0x44)


def _set_cell_bg(cell, hex_colour: str):
    """Fill a table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def _parse_inline(run_text: str, para):
    """
    Parse **bold** and `code` inline markers and add runs to para.
    """
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    last = 0
    for m in pattern.finditer(run_text):
        # Plain text before match
        if m.start() > last:
            para.add_run(run_text[last:m.start()])
        if m.group(2):                   # **bold**
            r = para.add_run(m.group(2))
            r.bold = True
        elif m.group(3):                 # `code`
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        last = m.end()
    if last < len(run_text):
        para.add_run(run_text[last:])


def _parse_md_table(doc: Document, lines: List[str]):
    """Render a markdown pipe-table as a Word table."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|?[-:| ]+\|?\s*$', line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= max_cols:
                break
            cell = row.cells[c_idx]
            # Strip markdown bold markers for header row
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            cell.text = clean
            if r_idx == 0:
                # Header row — purple background, white bold text
                _set_cell_bg(cell, '8B5CF6')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9)
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)


def convert(md_text: str, output_path: str):
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Default body style ──────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Headings ────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text  = re.sub(r'\*\*(.+?)\*\*', r'\1', h_match.group(2)).strip()
            # Strip emoji for cleaner docx headings
            text  = re.sub(r'[^\x00-\x7F\u2019\u2018\u201C\u201D]+\s*', '', text).strip() or text
            para  = doc.add_heading(text, level=min(level, 4))
            for run in para.runs:
                run.font.color.rgb = PURPLE
                run.font.bold = True
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if re.match(r'^---+$', line.strip()):
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pb   = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '8B5CF6')
            pb.append(bottom)
            pPr.append(pb)
            i += 1
            continue

        # ── Markdown table ───────────────────────────────────────────────────
        if re.match(r'^\s*\|', line):
            table_lines = []
            while i < len(lines) and re.match(r'^\s*\|', lines[i]):
                table_lines.append(lines[i])
                i += 1
            _parse_md_table(doc, table_lines)
            doc.add_paragraph()   # spacing after table
            continue

        # ── Bullet / list item ───────────────────────────────────────────────
        li_match = re.match(r'^(\s*)[*\-]\s+(.*)', line)
        if li_match:
            indent   = len(li_match.group(1)) // 2
            content  = li_match.group(2)
            style_name = 'List Bullet' if indent == 0 else 'List Bullet 2'
            para = doc.add_paragraph(style=style_name)
            _parse_inline(content, para)
            i += 1
            continue

        # ── Block quote / italic meta line ──────────────────────────────────
        bq_match = re.match(r'^>\s*(.*)', line)
        if bq_match:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run(bq_match.group(1))
            run.italic = True
            run.font.color.rgb = GREY
            run.font.size = Pt(9)
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Regular paragraph ────────────────────────────────────────────────
        para = doc.add_paragraph()
        _parse_inline(line, para)
        i += 1

    doc.save(output_path)
    print(f"Saved: {output_path}", file=sys.stderr)


def main():
    p = argparse.ArgumentParser(description='Convert Markdown report to .docx')
    p.add_argument('--input',  required=True, help='Input .md file path')
    p.add_argument('--output', required=True, help='Output .docx file path')
    args = p.parse_args()

    if not os.path.exists(args.input):
        print(f"ERROR: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    with open(args.input, 'r', encoding='utf-8') as f:
        md_text = f.read()

    convert(md_text, args.output)


if __name__ == '__main__':
    main()
